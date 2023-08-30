import streamlit as st
from docx import Document
import re
import os
from datetime import datetime
import tempfile
from io import BytesIO



"""

Copyrighted By Jayson Villena


"""
def replace_text_with_formatting(text, old_text, new_text):
    return re.sub(re.escape(old_text), new_text, text)

def replace_date_with_current_date(text, date_placeholder):
    current_date = datetime.now().strftime('%B %d, %Y')
    return re.sub(re.escape(date_placeholder), current_date, text)

def replace_date_and_word_in_docx(docx_path, old_word, new_word, date_placeholder):
    doc = Document(docx_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = replace_text_with_formatting(run.text, old_word, new_word)
            run.text = replace_date_with_current_date(run.text, date_placeholder)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = replace_text_with_formatting(run.text, old_word, new_word)
                        run.text = replace_date_with_current_date(run.text, date_placeholder)

    output_path = os.path.join(tempfile.gettempdir(), 'updated_' + os.path.basename(docx_path)+'.docx')
    doc.save(output_path)
    return output_path

st.title("DOCX Modifier and Downloader")

uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])
if uploaded_file is not None:
    old_word = st.text_input("Old Word to Replace")
    new_word = st.text_input("New Word")
    date_placeholder = st.text_input("Date Placeholder")

    if st.button("Modify and Download"):
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.read())
            temp_file_path = temp_file.name
        
        modified_path = replace_date_and_word_in_docx(temp_file_path, old_word, new_word, date_placeholder)
        
        st.success("DOCX file has been modified.")
        
        with open(modified_path, "rb") as file:
            file_bytes = file.read()
        
        st.download_button(
            label="Download Modified DOCX",
            data=BytesIO(file_bytes),
            file_name=os.path.basename(new_word + '_'+ uploaded_file.name),  # Use the original uploaded file name
            key="download_button")
    

"""

Copyrighted By Jayson Villena

"""
